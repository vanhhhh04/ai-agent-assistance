# -*- coding: utf-8 -*-
"""Thay 5 sơ đồ ASCII trong CH2 bằng hình Mermaid đã render; bỏ ghi chú ASCII cũ."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
import md_to_docx_thesis as T

BASE = Path(__file__).resolve().parent.parent / "documentations"
DIAG = BASE / "diagrams"
SRC = BASE / "CHUONG_2_PHAN_TICH_THIET_KE_THEM_BANG_FORMATTED.docx"
OUT = BASE / "CHUONG_2_PHAN_TICH_THIET_KE_DIAGRAMS.docx"

IMG = {
    "2.1": "hinh_2_1_usecase.png",
    "2.2": "hinh_2_2_kientruc.png",
    "2.3": "hinh_2_3_pipeline.png",
    "2.4": "hinh_2_4_starschema.png",
    "2.5": "hinh_2_5_sse.png",
}


def put_picture(par, img_path):
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    w, h = T.png_size(img_path) or (1, 1)
    if w >= h:
        run.add_picture(str(img_path), width=Cm(15.5))
    elif (h / w) * 13 > 21:
        run.add_picture(str(img_path), height=Cm(21))
    else:
        run.add_picture(str(img_path), width=Cm(13))


def main():
    d = Document(str(SRC))
    # bỏ dòng ghi chú ASCII cũ
    for p in list(d.paragraphs):
        if "Ghi chú cho tác giả" in p.text and "ASCII" in p.text:
            p._element.getparent().remove(p._element)

    ps = d.paragraphs
    done = 0
    for i, p in enumerate(ps):
        t = p.text.strip()
        if t.startswith("Hình 2."):
            num = t.split(":")[0].replace("Hình", "").strip()
            if num not in IMG:
                continue
            j = i - 1
            while j >= 0 and ps[j].text.strip() == "":
                j -= 1
            if j >= 0:
                put_picture(ps[j], DIAG / IMG[num])
                done += 1

    d.save(str(OUT))
    print(f"SAVED {OUT.name} | diagrams embedded: {done} | inline_images={len(d.inline_shapes)}")


if __name__ == "__main__":
    main()
