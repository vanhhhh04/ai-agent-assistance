# -*- coding: utf-8 -*-
"""Trích mục lục từ file đồ án đã ghép -> mucluc.docx (định dạng chuẩn)."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER

sys.path.insert(0, str(Path(__file__).resolve().parent))
import md_to_docx_thesis as T

BASE = Path(__file__).resolve().parent.parent / "documentations"
SRC = BASE / "DO_AN_TOT_NGHIEP_DataFinch_HOAN_CHINH.docx"
OUT = BASE / "mucluc.docx"

# Các mục phần đầu (FrontHeading) cần đưa vào mục lục (bỏ chính "MỤC LỤC")
FRONT_INCLUDE = {
    "DANH MỤC CÁC CHỮ VIẾT TẮT",
    "DANH MỤC BẢNG BIỂU",
    "DANH MỤC HÌNH VẼ",
}

def main():
    src = Document(str(SRC))
    entries = []  # (level, text)  level: 0 = mục lớn/H1, 1 = H2, 2 = H3
    for p in src.paragraphs:
        sn = p.style.name if p.style else ''
        t = p.text.strip()
        if not t:
            continue
        if sn == 'FrontHeading' and t.upper() in FRONT_INCLUDE:
            entries.append((0, t))
        elif sn == 'Heading 1':
            entries.append((0, t))
        elif sn == 'Heading 2':
            entries.append((1, t))
        elif sn == 'Heading 3':
            entries.append((2, t))

    doc = Document()
    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'; n.font.size = Pt(14)
    T._set_run_font_style(n, 'Times New Roman')
    s = doc.sections[0]
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0); s.left_margin = Cm(3.5); s.right_margin = Cm(2.0)

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    r = title.add_run("MỤC LỤC"); r.bold = True; r.font.size = Pt(16); T._set_run_font(r)

    RIGHT = Cm(15.5)
    for level, text in entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5; pf.space_after = Pt(2)
        pf.left_indent = Cm(0.8 * level)
        pf.tab_stops.add_tab_stop(RIGHT, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(text)
        run.bold = (level == 0)
        run.font.size = Pt(13 if level else 14)
        T._set_run_font(run)
        tab = p.add_run("\t"); T._set_run_font(tab)  # dấu chấm dẫn -> số trang (điền sau)

    try:
        doc.save(str(OUT))
        print("SAVED:", OUT.name, "| entries:", len(entries))
    except PermissionError:
        alt = OUT.with_name("mucluc_MOI.docx")
        doc.save(str(alt))
        print("FILE GOC DANG KHOA -> SAVED:", alt.name, "| entries:", len(entries))
    for lv, tx in entries:
        print('  ' + '  ' * lv + ('• ' if lv else '▸ ') + tx[:70])


if __name__ == '__main__':
    main()
