# -*- coding: utf-8 -*-
"""
Sinh CHUONG_3 / CHUONG_4 .docx theo CHUẨN FORMAT ĐỒ ÁN TỐT NGHIỆP từ file .md.

Quy định trình bày đã áp dụng:
  - Font Times New Roman 14, dãn dòng 1.5, canh đều (justify).
  - Lề: trên 2.5cm, dưới 2cm, trái 3.5cm, phải 2cm.
  - Số trang đánh GIỮA, phía TRÊN mỗi trang (header).
  - BẢNG: vẽ bảng Word thật (Table Grid); đầu đề "Bảng x.y:" đặt PHÍA TRÊN bảng.
  - HÌNH: ảnh/đoạn code; đầu đề "Hình x.y:" đặt PHÍA DƯỚI hình.
  - Đánh số bảng/hình gắn với số chương, tuần tự trong chương.
"""
import re
import struct
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent / "documentations"
SHOTS = BASE / "screenshots"
DIAGRAMS = BASE / "diagrams"
IMG_DIRS = [SHOTS, DIAGRAMS]

CODE_BG = "F2F2F2"
HEADER_BG = "D9E2F3"
NOTE_BG = "FFF2CC"
CONTENT_WIDTH_CM = 15.0  # 21 - 3.5 - 2 - lề an toàn

# ───────────────────────── cấu hình từng chương ─────────────────────────

# CH1 (bản rút gọn): 3 bảng (theo thứ tự xuất hiện) + 2 hình lý thuyết
CH1_TABLE_CAPTIONS = [
    "Một số mô hình ngôn ngữ lớn phổ biến",
    "So sánh Star schema và Snowflake schema",
    "So sánh BM25, tìm kiếm vector và truy hồi lai",
]
CH1_IMAGES = {
    "1.1": ["hinh_1_1_ai_ml_dl.png"],
    "1.2": ["hinh_1_3_rag.png"],
}

# CH3: bảng & hình (code) đánh số từ danh sách cố định
CH3_TABLE_CAPTIONS = [
    "Ngăn xếp công nghệ và phiên bản triển khai",
    "Ba bộ mô phỏng nguồn dữ liệu doanh nghiệp",
    "So sánh dialect rules giữa HiveQL và PostgreSQL",
    "Các nhóm service trong Docker Compose",
]
# CH2 (bản rút gọn): 9 bảng giữ lại (theo thứ tự) + 5 sơ đồ
CH2_TABLE_CAPTIONS = [
    "Các tác nhân (actors) của hệ thống",
    "Danh sách yêu cầu chức năng (Functional Requirements)",
    "Danh sách yêu cầu phi chức năng (Non-Functional Requirements)",
    "Ba tầng của kiến trúc Medallion",
    "Ba bảng fact và hạt (grain) tương ứng",
    "Các kiểm tra an toàn của Guardrails",
    "Hai backend thực thi và đặc điểm",
    "So sánh thiết kế OLTP (PostgreSQL) và OLAP (Hive Gold)",
    "Các endpoint REST API của Gateway",
]
CH2_IMAGES = {
    "2.1": ["hinh_2_1_usecase.png"],
    "2.2": ["hinh_2_2_kientruc.png"],
    "2.3": ["hinh_2_3_pipeline.png"],
    "2.4": ["hinh_2_4_starschema.png"],
    "2.5": ["hinh_2_5_sse.png"],
}

# CH3: danh sách hình HỢP NHẤT theo thứ tự xuất hiện trong tài liệu.
# Mỗi mục: (caption, files) — files=None là HÌNH MÃ NGUỒN (code fence);
# files=[...] là ẢNH CHỤP MÀN HÌNH (placeholder [Hình ...]). Số hình tự tăng tuần tự.
CH3_FIGURE_CAPTIONS = [
    ("Cấu trúc thư mục monorepo của dự án", None),                         # 3.1
    ("Vòng đời (lifespan) và khởi tạo ứng dụng FastAPI", None),            # 3.2
    ("Hàm _run() điều phối pipeline NL→SQL", None),                       # 3.3
    ("Giao diện trừu tượng LLMAdapter (Adapter Pattern)", None),          # 3.4
    ("Vòng retry với exponential backoff tại LLM Gateway", None),         # 3.5
    ("Bronze ingestion với Spark Structured Streaming (Trigger.AvailableNow)", None),  # 3.6
    ("Hàm latest_per_id — khử trùng CDC bằng window function", None),     # 3.7
    ("Dựng bảng fact_sales bằng JOIN star schema", None),                 # 3.8
    ("Airflow DAG medallion_pipeline (lịch */15 phút)", None),            # 3.9
    ("Truy vấn hybrid BM25 + kNN (bool/should) trên OpenSearch", None),   # 3.10
    ("Trích đoạn system prompt của Supervisor", None),                    # 3.11
    ("Schema augmentation chống ảo giác cột", None),                      # 3.12
    ("Hàm validate_sql — 7 lớp kiểm tra Guardrails", None),               # 3.13
    ("Giao diện đăng nhập và onboarding wizard 5 bước (demo)", ["02_login.png", "03_signup_step1.png"]),   # 3.14
    ("SSE client đọc luồng sự kiện ở frontend", None),                    # 3.15
    ("Giao diện chat NL→SQL với agent pipeline visualizer", ["05_ask_result.png"]),                        # 3.16
    ("Giao diện Saved queries lưu bằng localStorage (demo)", ["06_saved.png"]),                            # 3.17
    ("Giao diện Reports/Dashboard với Recharts (demo)", ["07_reports.png"]),                               # 3.18
    ("Giao diện quản lý Data Sources và Schema Catalog (demo)", ["08_data.png"]),                          # 3.19
    ("Giao diện Settings — provider-agnostic (demo)", ["09_settings.png"]),                                # 3.20
    ("Giao diện Billing và usage tracking (demo)", ["10_billing.png"]),                                    # 3.21
]

# CH4 (bản rút gọn): 7 bảng — mỗi phần thực nghiệm chỉ 1 bảng. (mô tả, số_liệu_minh_họa)
CH4_TABLES = [
    ("Phân bố bộ dữ liệu đánh giá (gold set) theo nhóm intent", False),
    ("So sánh độ chính xác giữa ba nhà cung cấp LLM", True),
    ("Ma trận nhầm lẫn phân loại intent", True),
    ("So sánh các phương pháp truy hồi (BM25, kNN, hybrid)", True),
    ("Độ chính xác SQL Writer theo phương ngữ (Claude Sonnet 4.6)", True),
    ("Hiệu quả 7 lớp kiểm tra của Guardrails", True),
    ("Độ trễ end-to-end theo phân vị (p50/p95/p99)", True),
]
# CH4 không còn ảnh giao diện (đã chuyển sang mục 3.7 của Chương 3)
CH4_IMAGES = {}

# ───────────────────────── tiện ích định dạng ─────────────────────────

def set_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _set_run_font(run, name='Times New Roman'):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(a), name)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*)')


def add_inline(paragraph, text):
    """Phân tích **đậm**, *nghiêng*, `mã`."""
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True; _set_run_font(r)
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); _set_run_font(r); r.font.size = Pt(13)  # inline code: TNR
        elif tok.startswith('*') and tok.endswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True; _set_run_font(r)
        else:
            r = paragraph.add_run(tok); _set_run_font(r)


def add_field(paragraph, instr):
    run = paragraph.add_run(); _set_run_font(run); run.font.size = Pt(13)
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '1'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(b); r.append(it); r.append(sep); r.append(t); r.append(e)


def png_size(path):
    try:
        with open(path, 'rb') as f:
            head = f.read(26)
        if head[:8] == b'\x89PNG\r\n\x1a\n':
            w, h = struct.unpack('>II', head[16:24])
            return w, h
    except Exception:
        pass
    return None


def add_table_caption(doc, chap, idx, desc, minhhoa=False):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6); cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = cap.add_run(f"Bảng {chap}.{idx}: {desc}")
    r.bold = True; r.font.size = Pt(13); _set_run_font(r)
    if minhhoa:
        r2 = cap.add_run("  (số liệu minh họa)"); r2.italic = True; r2.font.size = Pt(12); _set_run_font(r2)


def add_figure_caption(doc, chap, idx, desc):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = cap.add_run(f"Hình {chap}.{idx}: {desc}")
    r.bold = True; r.font.size = Pt(13); _set_run_font(r)


def add_code_block(doc, lines):
    p = doc.add_paragraph(); set_shading(p, CODE_BG)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, ln in enumerate(lines):
        run = p.add_run(ln); _set_run_font(run); run.font.size = Pt(10)  # code block: TNR, gọn
        if i < len(lines) - 1:
            run.add_break()


def _table_full_width(t):
    """Cho bảng chiếm trọn chiều rộng vùng nội dung (100%)."""
    t.autofit = True
    tblPr = t._tbl.tblPr
    for el in tblPr.findall(qn('w:tblW')):
        tblPr.remove(el)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'pct'); w.set(qn('w:w'), '5000')
    tblPr.append(w)


TABLE_FONT = 12  # cỡ chữ trong bảng; bảng "to hơn" chủ yếu nhờ full-width (xem _table_full_width)


def add_md_table(doc, rows):
    header, data = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _table_full_width(t)
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ''
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(c.paragraphs[0], cell)
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(TABLE_FONT)
        set_shading(c.paragraphs[0], HEADER_BG)
    for drow in data:
        cells = t.add_row().cells
        for j, cell in enumerate(drow):
            if j < len(cells):
                cells[j].paragraphs[0].text = ''
                add_inline(cells[j].paragraphs[0], cell)
                for run in cells[j].paragraphs[0].runs:
                    run.font.size = Pt(TABLE_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image_seq(doc, chap, idx, desc, files, width_cm=9.5):
    """Chèn ảnh theo danh sách file cụ thể (dùng cho CH3 — số hình tuần tự).
    width_cm mặc định 11.5 (nhỏ hơn lề) cho ảnh chụp màn hình landscape, tiết kiệm trang."""
    chosen = None
    for fn in files or []:
        for d in IMG_DIRS:
            if (d / fn).exists():
                chosen = d / fn; break
        if chosen:
            break
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    if chosen:
        size = png_size(chosen)
        run = p.add_run()
        if size and size[0] and (size[1] / size[0]) * width_cm > 20:
            run.add_picture(str(chosen), height=Cm(20))
        else:
            run.add_picture(str(chosen), width=Cm(width_cm))
    else:
        run = p.add_run("[ Ảnh chụp màn hình — bổ sung sau ]")
        run.italic = True; _set_run_font(run); set_shading(p, NOTE_BG)
    add_figure_caption(doc, chap, idx, desc)
    return chosen is not None


def add_image_figure(doc, num, desc, images_map=CH4_IMAGES):
    """Chèn ảnh (nếu có) hoặc placeholder, rồi caption 'Hình' phía dưới."""
    files = images_map.get(num, [])
    chosen = None
    for fn in files:
        for d in IMG_DIRS:
            if (d / fn).exists():
                chosen = d / fn; break
        if chosen:
            break
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    if chosen:
        size = png_size(chosen)
        run = p.add_run()
        if size and size[0] and (size[1] / size[0]) * CONTENT_WIDTH_CM > 20:
            run.add_picture(str(chosen), height=Cm(20))  # ảnh quá cao → giới hạn cao
        else:
            run.add_picture(str(chosen), width=Cm(CONTENT_WIDTH_CM))
    else:
        run = p.add_run("[ Ảnh chụp màn hình — bổ sung sau ]")
        run.italic = True; _set_run_font(run); set_shading(p, NOTE_BG)
    chap, idx = num.split('.')
    add_figure_caption(doc, chap, idx, desc)
    return chosen is not None


# ───────────────────────── thiết lập tài liệu ─────────────────────────

def build_document():
    doc = Document()
    # Normal: TNR 14, 1.5 line, justify
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'; normal.font.size = Pt(14)
    _set_run_font_style(normal, 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(6); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Heading styles → TNR đen, đậm
    for name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)]:
        st = doc.styles[name]
        st.font.name = 'Times New Roman'; st.font.size = Pt(size)
        st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
        _set_run_font_style(st, 'Times New Roman')
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.5
    # Lề + số trang trên-giữa
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.25)
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(hp, ' PAGE ')
    return doc


def _set_run_font_style(style, name):
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(a), name)


def parse_table_block(block):
    rows = []
    for ln in block:
        if set(ln.replace('|', '').strip()) <= set('-: '):
            continue
        rows.append([c.strip() for c in ln.strip().strip('|').split('|')])
    return rows


IMG_PLACEHOLDER_RE = re.compile(r'^`?\[Hình (\d+\.\d+):\s*([^\]]+)\]`?\s*(?:—\s*)?(.*)$')


def convert(chapter, src_name, out_name):
    src = BASE / src_name
    lines = src.read_text(encoding='utf-8').split('\n')
    doc = build_document()

    tbl_idx = fig_idx = img_done = 0
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # code fence → HÌNH (ch3) hoặc khối mã không số (ch4)
        if stripped.startswith('```'):
            code = []; i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            add_code_block(doc, code)
            if chapter == 3:
                fig_idx += 1
                entry = CH3_FIGURE_CAPTIONS[fig_idx - 1] if fig_idx - 1 < len(CH3_FIGURE_CAPTIONS) else ("Trích đoạn mã nguồn", None)
                add_figure_caption(doc, 3, fig_idx, entry[0])
            continue

        # bảng markdown → vẽ bảng + caption phía trên
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            tbl_idx += 1
            if chapter == 1:
                desc = CH1_TABLE_CAPTIONS[tbl_idx - 1] if tbl_idx - 1 < len(CH1_TABLE_CAPTIONS) else "Bảng dữ liệu"
                minh = False
            elif chapter == 2:
                desc = CH2_TABLE_CAPTIONS[tbl_idx - 1] if tbl_idx - 1 < len(CH2_TABLE_CAPTIONS) else "Bảng dữ liệu"
                minh = False
            elif chapter == 3:
                desc = CH3_TABLE_CAPTIONS[tbl_idx - 1] if tbl_idx - 1 < len(CH3_TABLE_CAPTIONS) else "Bảng dữ liệu"
                minh = False
            else:
                desc, minh = CH4_TABLES[tbl_idx - 1] if tbl_idx - 1 < len(CH4_TABLES) else ("Bảng dữ liệu", False)
            add_table_caption(doc, chapter, tbl_idx, desc, minh)
            add_md_table(doc, parse_table_block(block))
            continue

        # CH4: bỏ qua dòng caption cũ "**Bảng 4.x — ...**" (đã tự sinh caption)
        if chapter == 4 and re.match(r'^\*\*Bảng \d+\.\d+', stripped):
            i += 1; continue

        # CH3: placeholder ảnh [Hình x.y: ...] → chèn ảnh chụp màn hình, số hình tuần tự
        m = IMG_PLACEHOLDER_RE.match(stripped)
        if chapter == 3 and m:
            fig_idx += 1
            entry = CH3_FIGURE_CAPTIONS[fig_idx - 1] if fig_idx - 1 < len(CH3_FIGURE_CAPTIONS) else ("Ảnh giao diện", [])
            if add_image_seq(doc, 3, fig_idx, entry[0], entry[1] or []):
                img_done += 1
            i += 1
            continue

        # CH1/CH2/CH4: placeholder ảnh [Hình x.y: ...] → chèn ảnh + caption dưới + phần mô tả
        if chapter in (1, 2, 4) and m:
            num, cap, rest = m.group(1), m.group(2).strip(), m.group(3).strip()
            imap = {1: CH1_IMAGES, 2: CH2_IMAGES, 4: CH4_IMAGES}[chapter]
            if add_image_figure(doc, num, cap, imap):
                img_done += 1
            if rest:
                p = doc.add_paragraph(); add_inline(p, rest)
            i += 1; continue

        # heading
        if stripped.startswith('### '):
            doc.add_paragraph(stripped[4:].strip(), style='Heading 3')
        elif stripped.startswith('## '):
            doc.add_paragraph(stripped[3:].strip(), style='Heading 2')
        elif stripped.startswith('# '):
            h = doc.add_paragraph(stripped[2:].strip(), style='Heading 1')
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped == '---':
            pass
        elif stripped.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.0)
            set_shading(p, NOTE_BG)
            add_inline(p, stripped[2:].strip())
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number'); add_inline(p, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, stripped[2:])
        elif stripped == '':
            pass
        else:
            p = doc.add_paragraph(); add_inline(p, stripped)
        i += 1

    out = BASE / out_name
    doc.save(str(out))
    print(f"[CH{chapter}] saved: {out_name} | bảng={tbl_idx} hình_code={fig_idx if chapter==3 else 0} ảnh={img_done}")


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('--only', type=int, default=0, help='chỉ build 1 chương (1/3/4)')
    a = ap.parse_args()
    if a.only in (0, 1):
        convert(1, "CHUONG_1_CO_SO_LY_THUYET.md", "CHUONG_1_CO_SO_LY_THUYET.docx")
    if a.only in (0, 2):
        convert(2, "CHUONG_2_PHAN_TICH_THIET_KE.md", "CHUONG_2_PHAN_TICH_THIET_KE_DIAGRAMS.docx")
    if a.only in (0, 3):
        convert(3, "CHUONG_3_XAY_DUNG_TRIEN_KHAI.md", "CHUONG_3_XAY_DUNG_TRIEN_KHAI_DOAN.docx")
    if a.only in (0, 4):
        convert(4, "CHUONG_4_THUC_NGHIEM_DANH_GIA.md", "CHUONG_4_THUC_NGHIEM_DANH_GIA_DOAN.docx")
