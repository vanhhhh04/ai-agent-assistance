# -*- coding: utf-8 -*-
"""
Ghép 4 chương + phần đầu/cuối thành MỘT file đồ án tốt nghiệp hoàn chỉnh,
tuân thủ quy định trình bày (TNR 14, dãn dòng 1.5, lề 2.5/2/3.5/2, số trang
giữa-trên, đầu đề Bảng ở trên / Hình ở dưới, đánh số theo chương).

Bố cục đầu ra: Bìa → Mục lục (tự động) → Danh mục chữ viết tắt → Danh mục bảng
→ Danh mục hình → MỞ ĐẦU → Chương 1..4 → KẾT LUẬN → TÀI LIỆU THAM KHẢO → PHỤ LỤC.
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

sys.path.insert(0, str(Path(__file__).resolve().parent))
import md_to_docx_thesis as T  # tái dùng helper

BASE = Path(__file__).resolve().parent.parent / "documentations"
OUT = BASE / "DO_AN_TOT_NGHIEP_DataFinch_HOAN_CHINH.docx"

CH = {
    1: BASE / "CHUONG_1_CO_SO_LY_THUYET.docx",
    2: BASE / "CHUONG_2_PHAN_TICH_THIET_KE_DIAGRAMS.docx",
    3: BASE / "CHUONG_3_XAY_DUNG_TRIEN_KHAI_DOAN.docx",
    4: BASE / "CHUONG_4_THUC_NGHIEM_DANH_GIA_DOAN.docx",
}

THESIS_TITLE = ("NGHIÊN CỨU VÀ XÂY DỰNG HỆ THỐNG TRỢ LÝ DỮ LIỆU AI HỎI – ĐÁP "
                "BẰNG NGÔN NGỮ TỰ NHIÊN (NL→SQL) THEO KIẾN TRÚC ĐA TÁC TỬ "
                "KẾT HỢP NỀN TẢNG DỮ LIỆU MEDALLION")

# ───────────────────────── chuẩn hóa style/section ─────────────────────────

def apply_styles(doc, page_break_h1=True):
    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'; n.font.size = Pt(14)
    T._set_run_font_style(n, 'Times New Roman')
    pf = n.paragraph_format
    pf.line_spacing = 1.5; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_after = Pt(6); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13), ('Heading 4', 13)]:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = 'Times New Roman'; st.font.size = Pt(size)
            st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
            T._set_run_font_style(st, 'Times New Roman')
            st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.5
    # Header/Footer phải là TNR (mặc định Word là theme Calibri) — số trang ở header
    for name in ['Header', 'Footer']:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = 'Times New Roman'; st.font.size = Pt(13)
            T._set_run_font_style(st, 'Times New Roman')
    if page_break_h1 and 'Heading 1' in [s.name for s in doc.styles]:
        doc.styles['Heading 1'].paragraph_format.page_break_before = True


def fix_sections(doc, page_number=True, first_page_diff=False):
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.5); s.right_margin = Cm(2.0)
        s.header_distance = Cm(1.25)
        s.different_first_page_header_footer = first_page_diff
        if page_number:
            hp = s.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not hp.runs:
                T.add_field(hp, ' PAGE ')


def _set_margins(sec):
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.25)


def _clear_runs(paragraph):
    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)


def _page_header(sec):
    """Header giữa-trên có trường PAGE (số trang theo định dạng của section)."""
    sec.different_first_page_header_footer = False
    sec.header.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_runs(hp)
    T.add_field(hp, ' PAGE ')


def _empty_header(sec):
    """Header rỗng → trang không có số (dùng cho bìa)."""
    sec.different_first_page_header_footer = False
    sec.header.is_linked_to_previous = False
    _clear_runs(sec.header.paragraphs[0])


def set_pgnum(sec, fmt=None, start=None):
    """Đặt định dạng số trang cho section: fmt='lowerRoman'|'decimal', start=số bắt đầu."""
    sectPr = sec._sectPr
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pg = OxmlElement('w:pgNumType')
    if fmt:
        pg.set(qn('w:fmt'), fmt)
    if start is not None:
        pg.set(qn('w:start'), str(start))
    ref = sectPr.find(qn('w:pgMar'))
    if ref is not None:
        ref.addnext(pg)
    else:
        sectPr.append(pg)


def remap_ch2(doc):
    mapping = {'Title': 'Heading 1', 'Heading 1': 'Heading 2', 'Heading 2': 'Heading 3'}
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ''
        if sn in mapping:
            p.style = doc.styles[mapping[sn]]


def normalize(path, ch2=False):
    doc = Document(str(path))
    if ch2:
        remap_ch2(doc)
    apply_styles(doc); fix_sections(doc)
    return doc


import re as _re
_CAP_B = _re.compile(r'^Bảng \d+\.\d+')
_CAP_H = _re.compile(r'^Hình \d+\.\d+')

def collect_caps(doc):
    bang, hinh = [], []
    for p in doc.paragraphs:
        t = p.text.strip()
        if _CAP_B.match(t):
            bang.append(t)
        elif _CAP_H.match(t):
            hinh.append(t)
    return bang, hinh


# ───────────────────────── tiện ích phần đầu/cuối ─────────────────────────

def add_front_style(doc):
    names = [s.name for s in doc.styles]
    if 'FrontHeading' not in names:
        st = doc.styles.add_style('FrontHeading', WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles['Normal']
        st.font.name = 'Times New Roman'; st.font.size = Pt(15)
        st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
        T._set_run_font_style(st, 'Times New Roman')
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.page_break_before = True
        st.paragraph_format.space_after = Pt(14)
        st.paragraph_format.line_spacing = 1.5


def C(doc, text, bold=False, size=14, align='center', after=4, before=0, italic=False, caps_break=False):
    p = doc.add_paragraph()
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT, 'just': WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    T._set_run_font(r)
    return p


def heading1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _set_outline_level(paragraph, lvl):
    pPr = paragraph._p.get_or_add_pPr()
    for el in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(el)
    o = OxmlElement('w:outlineLvl'); o.set(qn('w:val'), str(lvl)); pPr.append(o)


def front_heading(doc, text, in_toc=False):
    p = doc.add_paragraph(text, style='FrontHeading')
    if in_toc:
        _set_outline_level(p, 0)  # đưa vào MỤC LỤC ở cấp 1 (qua TOC \u)
    return p


def body(doc, text, align='just', bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'just' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.0)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; T._set_run_font(r)
    T.add_inline(p, text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); T.add_inline(p, text); return p


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run(); r = run._r
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '(Bôi đen toàn bộ rồi nhấn F9 để cập nhật mục lục)'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for e in (fld, instr, sep, t, end):
        r.append(e)


def caption_list(doc, captions):
    for cap in captions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(cap); r.font.size = Pt(13); T._set_run_font(r)
        r2 = p.add_run('\t'); T._set_run_font(r2)


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    T.add_inline(p, text)


ABBREV = [
    ("AI", "Artificial Intelligence — Trí tuệ nhân tạo"),
    ("API", "Application Programming Interface — Giao diện lập trình ứng dụng"),
    ("BM25", "Best Matching 25 — Hàm xếp hạng tìm kiếm theo từ khóa"),
    ("CDC", "Change Data Capture — Bắt thay đổi dữ liệu"),
    ("CTE", "Common Table Expression — Biểu thức bảng chung trong SQL"),
    ("DAG", "Directed Acyclic Graph — Đồ thị có hướng không chu trình"),
    ("DLQ", "Dead Letter Queue — Hàng đợi bản ghi lỗi"),
    ("ERP", "Enterprise Resource Planning — Hoạch định nguồn lực doanh nghiệp"),
    ("ETL", "Extract – Transform – Load — Trích xuất, biến đổi, nạp dữ liệu"),
    ("EX", "Execution Accuracy — Độ chính xác thực thi"),
    ("HDFS", "Hadoop Distributed File System — Hệ thống tệp phân tán Hadoop"),
    ("HiveQL", "Hive Query Language — Ngôn ngữ truy vấn của Apache Hive"),
    ("kNN", "k-Nearest Neighbors — k láng giềng gần nhất"),
    ("KPI", "Key Performance Indicator — Chỉ số hiệu năng then chốt"),
    ("LLM", "Large Language Model — Mô hình ngôn ngữ lớn"),
    ("MRR", "Mean Reciprocal Rank — Hạng nghịch đảo trung bình"),
    ("NL→SQL", "Natural Language to SQL — Chuyển ngôn ngữ tự nhiên sang SQL"),
    ("OLAP", "Online Analytical Processing — Xử lý phân tích trực tuyến"),
    ("OLTP", "Online Transaction Processing — Xử lý giao dịch trực tuyến"),
    ("PII", "Personally Identifiable Information — Thông tin định danh cá nhân"),
    ("RBAC", "Role-Based Access Control — Kiểm soát truy cập theo vai trò"),
    ("SQL", "Structured Query Language — Ngôn ngữ truy vấn có cấu trúc"),
    ("SSE", "Server-Sent Events — Sự kiện do máy chủ đẩy"),
    ("VSR", "Valid-SQL Rate — Tỷ lệ SQL hợp lệ"),
]

# ───────────────────────── phần đầu (front matter) ─────────────────────────

def build_front(bang_caps, hinh_caps):
    doc = T.build_document()
    apply_styles(doc); add_front_style(doc)

    # ===== SECTION 0: BÌA — không đánh số trang =====
    _set_margins(doc.sections[0])
    _empty_header(doc.sections[0])

    # --- BÌA ---
    C(doc, "BỘ CÔNG THƯƠNG", bold=True, size=13, before=6, after=0)
    C(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP HÀ NỘI", bold=True, size=14, after=2)
    C(doc, "─────────────", size=12, after=24)
    C(doc, "ĐỒ ÁN TỐT NGHIỆP ĐẠI HỌC", bold=True, size=16, after=2)
    C(doc, "NGÀNH: CÔNG NGHỆ THÔNG TIN", bold=True, size=14, after=28)
    C(doc, "ĐỀ TÀI:", bold=True, size=13, after=4)
    C(doc, THESIS_TITLE, bold=True, size=15, after=30)
    C(doc, "Giảng viên hướng dẫn:  ……………………………………", size=14, align='center', after=4)
    C(doc, "Sinh viên thực hiện:  CAO VIỆT ANH", size=14, align='center', after=4)
    C(doc, "Mã số sinh viên:  ……………………", size=14, align='center', after=4)
    C(doc, "Lớp – Khóa:  ……………………", size=14, align='center', after=28)
    C(doc, "Hà Nội – 2026", bold=True, size=14, after=0)

    # ===== SECTION 1: LỜI CẢM ƠN + MỤC LỤC + DANH MỤC — số La Mã (i, ii, …) =====
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec1 = doc.sections[-1]; _set_margins(sec1); _page_header(sec1); set_pgnum(sec1, 'lowerRoman', 1)

    # --- LỜI CẢM ƠN ---
    _lc = front_heading(doc, "LỜI CẢM ƠN", in_toc=True); _lc.paragraph_format.page_break_before = False
    body(doc, "Lời đầu tiên, em xin gửi lời cảm ơn chân thành đến Ban Giám hiệu Trường Đại học Công nghiệp Hà Nội cùng toàn thể quý Thầy, Cô trong Khoa Công nghệ Thông tin đã tận tình giảng dạy, truyền đạt những kiến thức chuyên môn và tạo điều kiện thuận lợi cho em trong suốt quá trình học tập và nghiên cứu tại trường.")
    body(doc, "Em xin bày tỏ lòng biết ơn sâu sắc tới ThS. Lê Như Hiền – giảng viên hướng dẫn đã tận tình chỉ bảo, định hướng nghiên cứu và luôn dành thời gian góp ý cho em trong suốt quá trình thực hiện đồ án tốt nghiệp. Những ý kiến đóng góp quý báu của cô không chỉ giúp em hoàn thiện đề tài mà còn giúp em nâng cao tư duy nghiên cứu, khả năng giải quyết vấn đề và kỹ năng phát triển hệ thống phần mềm.")
    body(doc, "Trong quá trình thực hiện đề tài “Xây dựng hệ thống AI Agent hỗ trợ truy vấn dữ liệu doanh nghiệp bằng ngôn ngữ tự nhiên”, em đã có cơ hội vận dụng và tổng hợp nhiều kiến thức đã được học như phát triển hệ thống Backend, Data Engineering, Data Warehouse, xử lý dữ liệu lớn, Apache Kafka, Apache Spark, Apache Hive, OpenSearch, cùng các công nghệ Trí tuệ nhân tạo hiện đại như Large Language Models (LLMs), Retrieval-Augmented Generation (RAG) và kiến trúc Multi-Agent AI. Đây là cơ hội quý báu giúp em hiểu rõ hơn về quy trình xây dựng một hệ thống dữ liệu và AI hoàn chỉnh từ khâu thu thập, xử lý dữ liệu đến triển khai ứng dụng thực tế.")
    body(doc, "Bên cạnh đó, em xin gửi lời cảm ơn tới gia đình, bạn bè và những người đã luôn động viên, hỗ trợ và tạo điều kiện thuận lợi cho em trong suốt quá trình học tập cũng như thực hiện đồ án tốt nghiệp.")
    body(doc, "Mặc dù đã cố gắng hoàn thành đề tài với tất cả sự nỗ lực của bản thân, tuy nhiên do thời gian nghiên cứu, kinh nghiệm thực tế và kiến thức còn hạn chế nên đồ án khó tránh khỏi những thiếu sót nhất định. Em rất mong nhận được những ý kiến đóng góp, nhận xét từ quý Thầy, Cô để đề tài được hoàn thiện hơn.")
    C(doc, "Em xin trân trọng cảm ơn!", align='center', before=6, after=10)
    C(doc, "Sinh viên thực hiện", align='right', after=2)
    C(doc, "Cao Việt Anh", bold=True, align='right', after=0)

    # --- MỤC LỤC ---
    _ml = front_heading(doc, "MỤC LỤC")  # giữ page_break_before → sang trang mới
    add_toc_field(doc)

    # --- DANH MỤC CHỮ VIẾT TẮT ---
    front_heading(doc, "DANH MỤC CÁC CHỮ VIẾT TẮT", in_toc=True)
    tb = doc.add_table(rows=1, cols=2); tb.style = 'Table Grid'
    T._table_full_width(tb)
    h = tb.rows[0].cells
    for j, txt in enumerate(["Từ viết tắt", "Giải thích"]):
        h[j].paragraphs[0].text = ''
        rr = h[j].paragraphs[0].add_run(txt); rr.bold = True; rr.font.size = Pt(13); T._set_run_font(rr)
        T.set_shading(h[j].paragraphs[0], T.HEADER_BG)
    for ab, mean in ABBREV:
        cells = tb.add_row().cells
        cells[0].paragraphs[0].text = ''; r0 = cells[0].paragraphs[0].add_run(ab); r0.bold = True; r0.font.size = Pt(13); T._set_run_font(r0)
        cells[1].paragraphs[0].text = ''; r1 = cells[1].paragraphs[0].add_run(mean); r1.font.size = Pt(13); T._set_run_font(r1)

    # --- DANH MỤC BẢNG ---
    front_heading(doc, "DANH MỤC BẢNG BIỂU", in_toc=True)
    caption_list(doc, bang_caps)

    # --- DANH MỤC HÌNH ---
    front_heading(doc, "DANH MỤC HÌNH VẼ", in_toc=True)
    caption_list(doc, hinh_caps)

    # ===== SECTION 2: MỞ ĐẦU trở đi — số Ả Rập (1, 2, …) =====
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec2 = doc.sections[-1]; _set_margins(sec2); _page_header(sec2); set_pgnum(sec2, 'decimal', 1)

    # --- MỞ ĐẦU ---
    _md = heading1(doc, "MỞ ĐẦU"); _md.paragraph_format.page_break_before = False
    body(doc, "", bold_lead="1. Lý do chọn đề tài")
    body(doc, "Chuyển đổi số khiến khối lượng dữ liệu doanh nghiệp tăng nhanh, nhưng phần lớn dữ liệu vẫn nằm \"ngủ yên\" trong các hệ quản trị cơ sở dữ liệu vì rào cản kỹ thuật: muốn khai thác phải biết viết SQL và hiểu lược đồ dữ liệu. Các công cụ BI truyền thống đòi hỏi dựng sẵn báo cáo, kém linh hoạt với câu hỏi phát sinh. Sự phát triển của mô hình ngôn ngữ lớn (LLM) mở ra hướng đi mới: cho phép người dùng hỏi dữ liệu bằng ngôn ngữ tự nhiên (NL→SQL). Tuy nhiên, cách tiếp cận một-lời-nhắc (single-prompt) gặp vấn đề ảo giác (hallucination) và khó kiểm soát an toàn khi lược đồ lớn và câu hỏi bằng tiếng Việt. Từ thực tế đó, đề tài xây dựng hệ thống DataFinch — trợ lý dữ liệu AI theo kiến trúc đa tác tử (multi-agent) kết hợp nền tảng dữ liệu Medallion — nhằm giải quyết bài toán hỏi – đáp dữ liệu doanh nghiệp một cách chính xác, an toàn và thân thiện với người Việt.")
    body(doc, "", bold_lead="2. Mục đích nghiên cứu")
    body(doc, "Nghiên cứu, thiết kế và hiện thực hóa một hệ thống có khả năng: (i) tiếp nhận câu hỏi tiếng Việt và tự động sinh câu lệnh SQL đúng ngữ nghĩa; (ii) thực thi truy vấn trên kho dữ liệu phân tích và trả kết quả kèm giải thích, trực quan hóa; (iii) bảo đảm an toàn truy vấn và giảm thiểu ảo giác; (iv) chứng minh tính khả thi qua thực nghiệm định lượng.")
    body(doc, "", bold_lead="3. Đối tượng và phạm vi nghiên cứu")
    body(doc, "Đối tượng nghiên cứu là bài toán NL→SQL trong môi trường doanh nghiệp và kiến trúc đa tác tử cho hệ thống này. Phạm vi gồm: đường ống dữ liệu Medallion (Bronze → Silver → Gold) trên HDFS/Spark/Hive; lớp ngữ nghĩa truy hồi lai (BM25 + kNN) trên OpenSearch; cụm tác tử Supervisor – Metadata Retriever – SQL Writer – Guardrails – Executor trên nền FastAPI; và giao diện người dùng Next.js. Đề tài tập trung chứng minh tính khả thi của lõi NL→SQL; một số chức năng quản trị giao diện được dựng ở mức demo và nêu rõ hướng phát triển.")
    body(doc, "", bold_lead="4. Phương pháp nghiên cứu")
    body(doc, "Kết hợp nghiên cứu lý thuyết (khảo sát các công trình NL→SQL, kiến trúc multi-agent, kỹ thuật retrieval-augmented generation) với nghiên cứu thực nghiệm (xây dựng hệ thống chạy thật bằng Docker Compose, đo đạc theo bộ chỉ số định lượng: độ chính xác, độ trễ, tỷ lệ ảo giác và chi phí).")
    body(doc, "", bold_lead="5. Ý nghĩa khoa học và thực tiễn")
    body(doc, "Về khoa học, đề tài làm rõ giá trị của kiến trúc đa tác tử và lớp ngữ nghĩa song ngữ trong việc giảm ảo giác cho NL→SQL. Về thực tiễn, hệ thống cung cấp một công cụ hỏi – đáp dữ liệu cho doanh nghiệp Việt mà không yêu cầu người dùng biết SQL, có thể tái lập hoàn toàn trên một máy.")
    body(doc, "", bold_lead="6. Bố cục đồ án")
    body(doc, "Đồ án gồm bốn chương: Chương 1 — Tổng quan về phân tích dữ liệu doanh nghiệp và các công nghệ nền tảng; Chương 2 — Phân tích và thiết kế hệ thống; Chương 3 — Xây dựng và triển khai hệ thống; Chương 4 — Thực nghiệm và đánh giá hệ thống; cùng phần Mở đầu, Kết luận và Tài liệu tham khảo.")
    return doc


# ───────────────────────── phần cuối (back matter) ─────────────────────────

def build_back():
    doc = T.build_document()
    apply_styles(doc); add_front_style(doc)
    fix_sections(doc, page_number=True)

    heading1(doc, "KẾT LUẬN")
    body(doc, "Đồ án đã nghiên cứu và xây dựng thành công hệ thống DataFinch — trợ lý dữ liệu AI hỏi – đáp bằng ngôn ngữ tự nhiên theo kiến trúc đa tác tử kết hợp nền tảng dữ liệu Medallion, chạy thật trên một máy bằng Docker Compose.")
    body(doc, "", bold_lead="Các kết quả đạt được:")
    bullet(doc, "Xây dựng đường ống dữ liệu Medallion hoàn chỉnh (Bronze → Silver → Gold) với CDC thời gian thực, làm sạch – khử trùng dữ liệu và mô hình hóa star schema cho truy vấn phân tích.")
    bullet(doc, "Hiện thực cụm năm tác tử (Supervisor, Metadata Retriever, SQL Writer, Guardrails, Executor) trên FastAPI, hỗ trợ streaming SSE và kiến trúc độc lập nhà cung cấp LLM.")
    bullet(doc, "Đề xuất và áp dụng các kỹ thuật chống ảo giác (schema augmentation, dialect rules, truy hồi lai BM25 + kNN) và bảy lớp kiểm tra an toàn ở Guardrails.")
    bullet(doc, "Thực nghiệm định lượng cho thấy hệ thống đáp ứng các yêu cầu chức năng và phi chức năng đặt ra, với nút thắt hiệu năng nằm ở engine thực thi Hive — chỉ ra hướng tối ưu rõ ràng.")
    body(doc, "", bold_lead="Hạn chế:")
    body(doc, "Một số chức năng quản trị giao diện (xác thực thật, lưu trữ đa người dùng, báo cáo, thanh toán) mới ở mức demo; số liệu một số bảng đánh giá là minh họa và cần đo lại đầy đủ trên bộ benchmark quy mô lớn hơn.")
    body(doc, "", bold_lead="Hướng phát triển:")
    body(doc, "Tích hợp xác thực/phân quyền thật và lưu trữ phía máy chủ; thay engine thực thi (Trino/Spark SQL) để giảm độ trễ; mở rộng bộ benchmark và vòng lặp tự cải tiến từ nhật ký truy vấn; hoàn thiện các màn hình quản trị từ demo lên sản phẩm.")

    heading1(doc, "TÀI LIỆU THAM KHẢO")
    front = doc.add_paragraph(); rr = front.add_run("Tiếng Việt"); rr.bold = True; rr.italic = True; T._set_run_font(rr)
    add_reference(doc, "[1] Cao Việt Anh (2026), *Báo cáo đồ án tốt nghiệp: Hệ thống trợ lý dữ liệu AI DataFinch*, Trường Đại học Công nghiệp Hà Nội, Hà Nội.")
    p2 = doc.add_paragraph(); rr2 = p2.add_run("Tiếng Anh"); rr2.bold = True; rr2.italic = True; T._set_run_font(rr2)
    add_reference(doc, "[2] Vaswani A. và cộng sự (2017), \"Attention Is All You Need\", *Advances in Neural Information Processing Systems (NeurIPS)*, 30, 5998–6008.")
    add_reference(doc, "[3] Yu T. và cộng sự (2018), \"Spider: A Large-Scale Human-Labeled Dataset for Complex and Cross-Domain Semantic Parsing and Text-to-SQL Task\", *Proceedings of EMNLP 2018*, 3911–3921.")
    add_reference(doc, "[4] Reimers N., Gurevych I. (2019), \"Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks\", *Proceedings of EMNLP-IJCNLP 2019*, 3982–3992.")
    add_reference(doc, "[5] Lewis P. và cộng sự (2020), \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks\", *Advances in Neural Information Processing Systems (NeurIPS)*, 33, 9459–9474.")
    add_reference(doc, "[6] Robertson S., Zaragoza H. (2009), \"The Probabilistic Relevance Framework: BM25 and Beyond\", *Foundations and Trends in Information Retrieval*, 3(4), 333–389.")
    add_reference(doc, "[7] Uber Engineering (2024), *QueryGPT: Natural Language to SQL Using Generative AI*, https://www.uber.com/blog/query-gpt/.")
    add_reference(doc, "[8] Apache Software Foundation (2024), *Apache Spark, Apache Hive, Apache Kafka, Apache NiFi, Apache Airflow — Official Documentation*, https://apache.org.")
    add_reference(doc, "[9] Debezium Community (2024), *Debezium Documentation — Change Data Capture for PostgreSQL*, https://debezium.io/documentation/.")
    add_reference(doc, "[10] OpenSearch Project (2024), *OpenSearch Documentation — k-NN and Hybrid Search*, https://opensearch.org/docs/.")
    add_reference(doc, "[11] Anthropic (2025), *Claude API Documentation — Prompt Caching and Tool Use*, https://docs.anthropic.com.")

    heading1(doc, "PHỤ LỤC")
    front_heading_local(doc, "PHỤ LỤC A. Khung mã đánh giá (benchmark)")
    body(doc, "Khung mã gợi ý đặt tại cli/run_benchmark.py, gọi endpoint /api/query/ask, thu thập sự kiện SSE và chấm điểm theo bộ chỉ số ở mục 4.1 (xem chi tiết tại Chương 4).")
    body(doc, "", bold_lead="PHỤ LỤC B. Ảnh giao diện bổ sung")
    body(doc, "Một số ảnh chụp giao diện bổ sung (trang đăng nhập, trang chat trạng thái rỗng, các trang quản trị demo) được lưu trong thư mục documentations/screenshots của dự án.")
    return doc


def front_heading_local(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); T._set_run_font(r)
    return p


# ───────────────────────── ghép ─────────────────────────

def main():
    print("normalize chapters...")
    ch1 = normalize(CH[1])
    ch2 = normalize(CH[2], ch2=False)  # CH2 nay sinh tu generator chung (H1/H2/H3 dung) - khong remap
    ch3 = normalize(CH[3])
    ch4 = normalize(CH[4])

    bang, hinh = [], []
    for d in (ch1, ch2, ch3, ch4):
        b, h = collect_caps(d); bang += b; hinh += h
    print(f"  captions: bảng={len(bang)} hình={len(hinh)}")

    front = build_front(bang, hinh)
    back = build_back()

    print("compose...")
    comp = Composer(front)
    for d in (ch1, ch2, ch3, ch4, back):
        comp.append(d)
    try:
        comp.save(str(OUT))
        print("SAVED:", OUT.name)
    except PermissionError:
        alt = OUT.with_name("DO_AN_TOT_NGHIEP_DataFinch_HOAN_CHINH_MOI.docx")
        comp.save(str(alt))
        print("FILE GOC DANG MO TRONG WORD -> SAVED:", alt.name)


if __name__ == '__main__':
    main()
