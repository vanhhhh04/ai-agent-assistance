# -*- coding: utf-8 -*-
"""
Chuyển CHUONG_2_PHAN_TICH_THIET_KE.md -> .docx cho bản đồ án.
Giữ nguyên nội dung & cấu trúc; CHỈ THÊM:
  - 18 bảng markdown được vẽ thành bảng Word thật (style 'Table Grid', có khung)
    kèm caption "Bảng 2.x: ..." phía trên.
  - Caption "Hình 2.x: ..." phía dưới mỗi sơ đồ/đoạn code (giữ nguyên dạng text).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent / "documentations"
SRC = BASE / "CHUONG_2_PHAN_TICH_THIET_KE.md"
OUT = BASE / "CHUONG_2_PHAN_TICH_THIET_KE_THEM_BANG.docx"

TABLE_CAPTIONS = [
    "Bảng 2.1: Các tác nhân (actors) của hệ thống",
    "Bảng 2.2: Đặc tả ca sử dụng UC1 — Hỏi dữ liệu bằng ngôn ngữ tự nhiên",
    "Bảng 2.3: Danh sách yêu cầu chức năng (Functional Requirements)",
    "Bảng 2.4: Danh sách yêu cầu phi chức năng (Non-Functional Requirements)",
    "Bảng 2.5: Ba tầng của kiến trúc Medallion",
    "Bảng 2.6: Ba nguồn dữ liệu và cơ chế tiếp nhận",
    "Bảng 2.7: Ba bảng fact và hạt (grain) tương ứng",
    "Bảng 2.8: Thiết kế đầu vào/đầu ra của Supervisor Agent",
    "Bảng 2.9: Ba index ngữ nghĩa và Top-K mặc định",
    "Bảng 2.10: Năm lớp chống ảo giác của SQL Writer",
    "Bảng 2.11: Các kiểm tra an toàn của Guardrails",
    "Bảng 2.12: Hai backend thực thi và đặc điểm",
    "Bảng 2.13: Lược đồ cơ sở dữ liệu ERP nguồn (PostgreSQL)",
    "Bảng 2.14: Lược đồ bảng gold.fact_sales (rút gọn)",
    "Bảng 2.15: So sánh thiết kế OLTP (PostgreSQL) và OLAP (Hive Gold)",
    "Bảng 2.16: Các trường của document cột trong finch_catalog",
    "Bảng 2.17: Các trang quản lý của ứng dụng",
    "Bảng 2.18: Các endpoint REST API của Gateway",
]
FIGURE_CAPTIONS = [
    "Hình 2.1: Sơ đồ ca sử dụng (use case) của hệ thống DataFinch",
    "Hình 2.2: Kiến trúc tổng thể hai trụ cột (Data Pipeline + AI Agent Service)",
    "Hình 2.3: Quy trình xử lý truy vấn NL→SQL tám bước",
    "Hình 2.4: Cấu trúc star schema của tầng Gold",
    "Hình 2.5: Ví dụ luồng sự kiện Server-Sent Events (SSE)",
]

CODE_BG = "F2F2F2"


def set_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`)')


def add_inline(paragraph, text):
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            paragraph.add_run(tok)


def add_code_block(doc, lines, caption):
    p = doc.add_paragraph(); set_shading(p, CODE_BG)
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.left_indent = Pt(6)
    for i, ln in enumerate(lines):
        run = p.add_run(ln); run.font.name = 'Consolas'; run.font.size = Pt(9)
        if i < len(lines) - 1:
            run.add_break()
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption); cr.italic = True; cr.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(10)


def add_md_table(doc, rows, caption):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption); cr.bold = True; cr.font.size = Pt(11)
    header, data = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ''
        add_inline(c.paragraphs[0], cell)
        for run in c.paragraphs[0].runs:
            run.bold = True
        set_shading(c.paragraphs[0], "D9E2F3")
    for drow in data:
        cells = t.add_row().cells
        for j, cell in enumerate(drow):
            if j < len(cells):
                cells[j].paragraphs[0].text = ''
                add_inline(cells[j].paragraphs[0], cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_table_block(block):
    rows = []
    for ln in block:
        if set(ln.replace('|', '').strip()) <= set('-: '):
            continue
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def main():
    lines = SRC.read_text(encoding='utf-8').split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(13)

    tbl_idx = fig_idx = 0
    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if stripped.startswith('```'):
            code = []; i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            cap = FIGURE_CAPTIONS[fig_idx] if fig_idx < len(FIGURE_CAPTIONS) else f"Hình 2.{fig_idx+1}"
            fig_idx += 1; add_code_block(doc, code, cap); continue
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            rows = parse_table_block(block)
            cap = TABLE_CAPTIONS[tbl_idx] if tbl_idx < len(TABLE_CAPTIONS) else f"Bảng 2.{tbl_idx+1}"
            tbl_idx += 1; add_md_table(doc, rows, cap); continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped == '---':
            pass
        elif stripped.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(stripped[2:].strip().strip('*')); r.italic = True
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number'); add_inline(p, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, stripped[2:])
        elif stripped == '':
            pass
        else:
            p = doc.add_paragraph(); add_inline(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Tables drawn: {tbl_idx} | Figures captioned: {fig_idx}")


if __name__ == '__main__':
    main()
