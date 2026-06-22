# -*- coding: utf-8 -*-
"""
Chuyển CHUONG_4_THUC_NGHIEM_DANH_GIA.md -> .docx cho bản đồ án.
Giữ nguyên nội dung & cấu trúc; CHỈ THÊM: vẽ 13 bảng markdown thành bảng Word
thật (style 'Table Grid', có khung).

Khác Ch2/Ch3: Chương 4 đã có sẵn caption "Bảng 4.1–4.9" viết trong text và đã
dùng "Hình 4.1–4.7" cho ảnh UI (mục 4.8). Do đó KHÔNG chèn thêm số caption nào
(tránh đụng numbering có sẵn); code/công thức giữ nguyên dạng text (không thêm Hình).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent / "documentations"
SRC = BASE / "CHUONG_4_THUC_NGHIEM_DANH_GIA.md"
OUT = BASE / "CHUONG_4_THUC_NGHIEM_DANH_GIA_THEM_BANG.docx"

CODE_BG = "F2F2F2"
INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`)')


def set_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_inline(paragraph, text, italic=False):
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            r = paragraph.add_run(tok)
        if italic:
            r.italic = True


def add_code_block(doc, lines):
    p = doc.add_paragraph(); set_shading(p, CODE_BG)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = Pt(6)
    for i, ln in enumerate(lines):
        run = p.add_run(ln); run.font.name = 'Consolas'; run.font.size = Pt(9)
        if i < len(lines) - 1:
            run.add_break()


def add_md_table(doc, rows):
    header, data = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ''
        add_inline(c.paragraphs[0], cell)
        for run in c.paragraphs[0].runs:
            run.bold = True
        set_shading(c.paragraphs[0], "D9E2F3")
    for drow in data:
        cells = t.add_row().cells
        for j, cell in enumerate(drow):
            if j < len(cells):
                cells[j].paragraphs[0].text = ''
                add_inline(cells[j].paragraphs[0], cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_table_block(block):
    rows = []
    for ln in block:
        if set(ln.replace('|', '').strip()) <= set('-: '):
            continue
        rows.append([c.strip() for c in ln.strip().strip('|').split('|')])
    return rows


def main():
    lines = SRC.read_text(encoding='utf-8').split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(13)

    tbl_idx = code_idx = 0
    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if stripped.startswith('```'):
            code = []; i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1; code_idx += 1; add_code_block(doc, code); continue
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            tbl_idx += 1; add_md_table(doc, parse_table_block(block)); continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped == '---':
            pass
        elif stripped.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            set_shading(p, "FFF2CC")  # nền vàng nhạt cho khối ghi chú
            add_inline(p, stripped[2:].strip(), italic=True)
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number'); add_inline(p, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, stripped[2:])
        elif stripped == '':
            pass
        else:
            p = doc.add_paragraph(); add_inline(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Tables drawn: {tbl_idx} | Code blocks (kept as text): {code_idx}")


if __name__ == '__main__':
    main()
