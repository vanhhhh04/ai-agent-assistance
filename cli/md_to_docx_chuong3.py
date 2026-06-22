# -*- coding: utf-8 -*-
"""
Chuyển CHUONG_3_XAY_DUNG_TRIEN_KHAI.md -> .docx cho bản đồ án.
Giữ nguyên nội dung & cấu trúc; CHỈ THÊM:
  - 4 bảng markdown được vẽ thành bảng Word thật (style 'Table Grid', có khung)
    kèm caption "Bảng 3.x: ..." phía trên.
  - Caption "Hình 3.x: ..." phía dưới mỗi đoạn code (code giữ nguyên dạng text).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent / "documentations"
SRC = BASE / "CHUONG_3_XAY_DUNG_TRIEN_KHAI.md"
OUT = BASE / "CHUONG_3_XAY_DUNG_TRIEN_KHAI_THEM_BANG.docx"

# Caption cho 4 bảng (theo thứ tự xuất hiện)
TABLE_CAPTIONS = [
    "Bảng 3.1: Ngăn xếp công nghệ và phiên bản triển khai",
    "Bảng 3.2: Ba bộ mô phỏng nguồn dữ liệu doanh nghiệp",
    "Bảng 3.3: So sánh dialect rules giữa HiveQL và PostgreSQL",
    "Bảng 3.4: Các nhóm service trong Docker Compose",
]
# Caption cho 19 đoạn code (theo thứ tự xuất hiện)
FIGURE_CAPTIONS = [
    "Hình 3.1: Cấu trúc thư mục monorepo của dự án",
    "Hình 3.2: Vòng đời (lifespan) và khởi tạo ứng dụng FastAPI",
    "Hình 3.3: Endpoint /api/query/ask trả về StreamingResponse (SSE)",
    "Hình 3.4: Hàm _run() điều phối pipeline NL→SQL",
    "Hình 3.5: Giao diện trừu tượng LLMAdapter (Adapter Pattern)",
    "Hình 3.6: Registry pattern và hàm get_adapter() (lazy import)",
    "Hình 3.7: Anthropic adapter — prompt caching và adaptive thinking",
    "Hình 3.8: Vòng retry với exponential backoff tại LLM Gateway",
    "Hình 3.9: Bronze ingestion với Spark Structured Streaming (Trigger.AvailableNow)",
    "Hình 3.10: Hàm latest_per_id — khử trùng CDC bằng window function",
    "Hình 3.11: Dựng bảng fact_sales bằng JOIN star schema",
    "Hình 3.12: Hàm write_gold — đăng ký EXTERNAL table vào Hive Metastore",
    "Hình 3.13: Airflow DAG medallion_pipeline (lịch */15 phút)",
    "Hình 3.14: Truy vấn hybrid BM25 + kNN (bool/should) trên OpenSearch",
    "Hình 3.15: Trích đoạn system prompt của Supervisor",
    "Hình 3.16: Schema augmentation chống ảo giác cột",
    "Hình 3.17: Hàm validate_sql — 7 lớp kiểm tra Guardrails",
    "Hình 3.18: Tổ chức route groups của frontend Next.js",
    "Hình 3.19: SSE client đọc luồng sự kiện ở frontend",
]

# Đoạn kết bổ sung của mục 3.9 (có trong file Word đính kèm, không có trong .md)
EXTRA_39 = ("Một số hạng mục được nêu trung thực là **chưa phát triển đầy đủ** "
            "(xác thực/phân quyền thật, persistence đa người dùng cho saved "
            "queries/reports) — đây là các hướng mở rộng, không ảnh hưởng tới việc "
            "chứng minh tính khả thi của lõi hệ thống. Trên cơ sở hệ thống đã xây "
            "dựng, **Chương 4** sẽ tiến hành thực nghiệm và đánh giá: so sánh các nhà "
            "cung cấp LLM, đo độ chính xác/độ trễ/tỷ lệ ảo giác của từng module và "
            "toàn pipeline, cùng trình diễn giao diện.")

CODE_BG = "F2F2F2"


def set_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`)')


def add_inline(paragraph, text):
    """Phân tích **bold** và `code` inline."""
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            paragraph.add_run(tok)


def add_code_block(doc, lines, caption):
    p = doc.add_paragraph()
    set_shading(p, CODE_BG)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(6)
    for i, ln in enumerate(lines):
        run = p.add_run(ln)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        if i < len(lines) - 1:
            run.add_break()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption); cr.italic = True; cr.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(10)


def add_md_table(doc, rows, caption):
    # caption phía trên
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption); cr.bold = True; cr.font.size = Pt(11)
    header, data = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].text = ''
        add_inline(c.paragraphs[0], cell)
        for run in c.paragraphs[0].runs:
            run.bold = True
        set_shading(c.paragraphs[0], "D9E2F3")
    for drow in data:
        cells = t.add_row().cells
        for j, cell in enumerate(drow):
            if j < len(cells):
                cells[j].paragraphs[0].text = ''
                add_inline(cells[j].paragraphs[0], cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def parse_table_block(block):
    rows = []
    for ln in block:
        if re.match(r'^\s*\|?\s*[-:\s|]+\s*$', ln) and set(ln.replace('|', '').strip()) <= set('-: '):
            continue  # dòng phân cách |---|
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def main():
    text = SRC.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(13)

    tbl_idx = 0
    fig_idx = 0
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- code fence ---
        if stripped.startswith('```'):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i])
                i += 1
            i += 1  # bỏ qua fence đóng
            cap = FIGURE_CAPTIONS[fig_idx] if fig_idx < len(FIGURE_CAPTIONS) else f"Hình 3.{fig_idx+1}"
            fig_idx += 1
            add_code_block(doc, code, cap)
            continue

        # --- bảng markdown ---
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i])
                i += 1
            rows = parse_table_block(block)
            cap = TABLE_CAPTIONS[tbl_idx] if tbl_idx < len(TABLE_CAPTIONS) else f"Bảng 3.{tbl_idx+1}"
            tbl_idx += 1
            add_md_table(doc, rows, cap)
            continue

        # --- heading ---
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:].strip(), level=0)
        elif stripped == '---':
            pass  # bỏ qua đường kẻ ngang
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(stripped[2:].strip().strip('*')); r.italic = True
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:])
        elif stripped == '':
            pass
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1

    # đoạn kết bổ sung 3.9
    p = doc.add_paragraph()
    add_inline(p, EXTRA_39)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Tables drawn: {tbl_idx} | Figures captioned: {fig_idx}")


if __name__ == '__main__':
    main()
